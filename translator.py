#!/usr/bin/env python3
"""
Система качественного перевода на основе MQM Framework и LQA методологии
"""
import os
import json
from docx import Document
from anthropic import Anthropic
from typing import Literal, Dict, Any
from dataclasses import dataclass, asdict


@dataclass
class TranslationResult:
    """Результат перевода с метаданными качества"""
    original_text: str
    translated_text: str
    mode: str
    source_lang: str
    target_lang: str
    passes: Dict[str, Any]
    final_score: Dict[str, str]


class QualityTranslator:
    """Переводчик с многоэтапной проверкой качества"""
    
    def __init__(self, api_key: str = None):
        self.client = Anthropic(api_key=api_key or os.environ.get("ANTHROPIC_API_KEY"))
        self.model = "claude-3-7-sonnet-20250219"
    
    def translate(
        self,
        text: str,
        source_lang: str,
        target_lang: str,
        mode: Literal["literary", "precise"] = "precise"
    ) -> TranslationResult:
        """
        Выполняет многоэтапный перевод с проверкой качества
        
        Args:
            text: Исходный текст
            source_lang: Язык оригинала
            target_lang: Целевой язык
            mode: "literary" (художественный) или "precise" (точный)
        """
        passes = {}
        
        # Pass 1: Начальный перевод
        print("🔄 Pass 1: Начальный перевод...")
        initial = self._initial_translation(text, source_lang, target_lang, mode)
        passes["initial"] = initial
        
        # Pass 2: Проверка орфографии и грамматики
        print("✍️  Pass 2: Проверка орфографии и грамматики...")
        spelling = self._check_spelling(initial, target_lang)
        passes["spelling"] = spelling
        
        # Pass 3: Проверка естественности
        print("🗣️  Pass 3: Проверка естественности...")
        fluency = self._check_fluency(spelling["corrected"], target_lang)
        passes["fluency"] = fluency
        
        # Pass 4: Сравнение смысла
        print("🔍 Pass 4: Сравнение смысла...")
        accuracy = self._check_accuracy(text, fluency["improved"], source_lang, target_lang)
        passes["accuracy"] = accuracy
        
        # Pass 5: Сравнение эмоционального тона
        print("🎭 Pass 5: Сравнение эмоционального тона...")
        tone = self._check_tone(text, accuracy["corrected"], source_lang, target_lang, mode)
        passes["tone"] = tone
        
        # Pass 6: Финальная проверка и полировка
        print("✨ Pass 6: Финализация...")
        final = self._finalize(tone["corrected"], target_lang, mode)
        passes["final"] = final
        
        # Итоговая оценка
        score = self._calculate_score(passes)
        
        return TranslationResult(
            original_text=text,
            translated_text=final["text"],
            mode=mode,
            source_lang=source_lang,
            target_lang=target_lang,
            passes=passes,
            final_score=score
        )
    
    def _call_claude(self, prompt: str, system: str = None) -> str:
        """Вызов Claude API"""
        messages = [{"role": "user", "content": prompt}]
        kwargs = {"model": self.model, "max_tokens": 8000, "messages": messages}
        if system:
            kwargs["system"] = system
        response = self.client.messages.create(**kwargs)
        return response.content[0].text
    
    def _initial_translation(self, text: str, source: str, target: str, mode: str) -> str:
        """Pass 1: Начальный перевод"""
        mode_desc = {
            "literary": "художественный перевод: передай смысл и тон максимально естественно и красиво",
            "precise": "точный перевод: передай смысл с максимальной точностью, используя правильные термины"
        }
        
        prompt = f"""Переведи текст с {source} на {target}.
Режим: {mode_desc[mode]}

Текст:
{text}

Верни только перевод без комментариев."""
        
        return self._call_claude(prompt)
    
    def _check_spelling(self, text: str, lang: str) -> Dict[str, Any]:
        """Pass 2: Проверка орфографии и грамматики"""
        prompt = f"""Проверь орфографию и грамматику текста на {lang}.

Текст:
{text}

Верни JSON:
{{
  "errors": ["список найденных ошибок"],
  "corrected": "исправленный текст"
}}"""
        
        response = self._call_claude(prompt)
        return json.loads(response.strip("```json\n").strip("```"))
    
    def _check_fluency(self, text: str, lang: str) -> Dict[str, Any]:
        """Pass 3: Проверка естественности"""
        prompt = f"""Проверь естественность текста на {lang}. Звучит ли он как написанный носителем языка?

Текст:
{text}

Верни JSON:
{{
  "issues": ["список неестественных конструкций"],
  "improved": "улучшенный текст"
}}"""
        
        response = self._call_claude(prompt)
        return json.loads(response.strip("```json\n").strip("```"))
    
    def _check_accuracy(self, original: str, translated: str, source: str, target: str) -> Dict[str, Any]:
        """Pass 4: Проверка точности передачи смысла"""
        prompt = f"""Сравни оригинал ({source}) и перевод ({target}). Совпадает ли смысл?

Оригинал:
{original}

Перевод:
{translated}

Верни JSON:
{{
  "matches": true/false,
  "discrepancies": ["список расхождений в смысле"],
  "corrected": "исправленный перевод (если нужно)"
}}"""
        
        response = self._call_claude(prompt)
        return json.loads(response.strip("```json\n").strip("```"))
    
    def _check_tone(self, original: str, translated: str, source: str, target: str, mode: str) -> Dict[str, Any]:
        """Pass 5: Проверка эмоционального тона"""
        prompt = f"""Сравни эмоциональный тон оригинала ({source}) и перевода ({target}).

Оригинал:
{original}

Перевод:
{translated}

Верни JSON:
{{
  "tone_matches": true/false,
  "original_tone": "описание тона оригинала",
  "translated_tone": "описание тона перевода",
  "corrected": "исправленный перевод (если нужно)"
}}"""
        
        response = self._call_claude(prompt)
        return json.loads(response.strip("```json\n").strip("```"))
    
    def _finalize(self, text: str, lang: str, mode: str) -> Dict[str, str]:
        """Pass 6: Финальная полировка"""
        prompt = f"""Финальная полировка перевода на {lang} (режим: {mode}).
Убедись что текст идеален.

Текст:
{text}

Верни JSON:
{{
  "text": "финальный отполированный текст"
}}"""
        
        response = self._call_claude(prompt)
        return json.loads(response.strip("```json\n").strip("```"))
    
    def _calculate_score(self, passes: Dict) -> Dict[str, str]:
        """Расчет итоговой оценки по MQM категориям"""
        return {
            "accuracy": "✓" if passes["accuracy"].get("matches", True) else "⚠",
            "fluency": "✓" if len(passes["fluency"].get("issues", [])) == 0 else "⚠",
            "spelling": "✓" if len(passes["spelling"].get("errors", [])) == 0 else "⚠",
            "tone": "✓" if passes["tone"].get("tone_matches", True) else "⚠"
        }


def extract_text_from_docx(filepath: str) -> str:
    """Извлекает текст из .docx файла"""
    doc = Document(filepath)
    return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def save_translation(result: TranslationResult, output_path: str):
    """Сохраняет перевод в .docx файл"""
    doc = Document()
    doc.add_heading(f'Перевод ({result.mode})', 0)
    doc.add_heading(f'{result.source_lang} → {result.target_lang}', level=2)
    
    # Добавляем переведенный текст
    for para in result.translated_text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    # Добавляем метаданные качества
    doc.add_page_break()
    doc.add_heading('Отчет о качестве', level=1)
    doc.add_paragraph(f"Точность (Accuracy): {result.final_score['accuracy']}")
    doc.add_paragraph(f"Беглость (Fluency): {result.final_score['fluency']}")
    doc.add_paragraph(f"Орфография (Spelling): {result.final_score['spelling']}")
    doc.add_paragraph(f"Тон (Tone): {result.final_score['tone']}")
    
    doc.save(output_path)


if __name__ == "__main__":
    # Пример использования
    translator = QualityTranslator()
    
    # Тестовый перевод
    test_text = "Hello, this is a test translation system."
    result = translator.translate(test_text, "English", "Russian", mode="precise")
    
    print("\n" + "="*60)
    print("РЕЗУЛЬТАТ ПЕРЕВОДА")
    print("="*60)
    print(f"\nОригинал: {result.original_text}")
    print(f"\nПеревод: {result.translated_text}")
    print(f"\nОценка качества:")
    for category, score in result.final_score.items():
        print(f"  {category}: {score}")
