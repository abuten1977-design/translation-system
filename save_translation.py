#!/usr/bin/env python3
"""
Сохранение результата перевода из JSON в .docx
"""
import argparse
import json
from docx import Document


def save_translation(json_path: str, output_path: str):
    """Сохраняет перевод из JSON в .docx"""
    # Загрузка JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Создание документа
    doc = Document()
    doc.add_heading('Перевод', 0)
    
    # Добавляем переведенный текст
    translated = data.get('translated_text', '')
    for para in translated.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    # Добавляем отчет о качестве
    if 'quality_report' in data:
        doc.add_page_break()
        doc.add_heading('Отчет о качестве (MQM)', level=1)
        
        report = data['quality_report']
        doc.add_paragraph(f"✓ Точность (Accuracy): {report.get('accuracy', 'N/A')}")
        doc.add_paragraph(f"✓ Естественность (Fluency): {report.get('fluency', 'N/A')}")
        doc.add_paragraph(f"✓ Орфография (Spelling): {report.get('spelling', 'N/A')}")
        doc.add_paragraph(f"✓ Тон (Tone): {report.get('tone', 'N/A')}")
    
    # Добавляем заметки
    if 'notes' in data:
        doc.add_paragraph()
        doc.add_heading('Заметки', level=2)
        doc.add_paragraph(data['notes'])
    
    doc.save(output_path)
    print(f"✅ Перевод сохранен в {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Сохранить перевод из JSON в .docx")
    parser.add_argument("json_file", help="JSON файл с результатом перевода")
    parser.add_argument("output", help="Выходной .docx файл")
    
    args = parser.parse_args()
    save_translation(args.json_file, args.output)


if __name__ == "__main__":
    main()
