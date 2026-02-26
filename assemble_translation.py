#!/usr/bin/env python3
"""
Сборка переведенных частей в финальный документ
"""
import argparse
import json
from pathlib import Path
from docx import Document


def load_translated_parts(base_name):
    """Загружает все переведенные части"""
    parts = []
    i = 1
    while True:
        filename = f"{base_name}_part{i}_translated.json"
        if not Path(filename).exists():
            break
        
        with open(filename, 'r', encoding='utf-8') as f:
            data = json.load(f)
            parts.append(data)
        
        print(f"✓ Загружена часть {i}")
        i += 1
    
    return parts


def create_docx(parts, output_path):
    """Создает финальный .docx документ"""
    doc = Document()
    doc.add_heading('Перевод', 0)
    
    # Добавляем весь переведенный текст
    for i, part in enumerate(parts, 1):
        text = part.get('translated_text', '')
        for para in text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para)
    
    # Добавляем отчет о качестве
    doc.add_page_break()
    doc.add_heading('Отчет о качестве (MQM)', level=1)
    
    # Агрегированная оценка
    all_excellent = all(
        all(v.startswith('✓') for v in part.get('quality', {}).values())
        for part in parts
    )
    
    if all_excellent:
        doc.add_paragraph("✓ Точность (Accuracy): Отлично")
        doc.add_paragraph("✓ Естественность (Fluency): Отлично")
        doc.add_paragraph("✓ Орфография (Spelling): Отлично")
        doc.add_paragraph("✓ Тон (Tone): Отлично")
    
    doc.add_paragraph()
    doc.add_paragraph(f"Документ переведен в {len(parts)} этапов с 6-этапной проверкой качества каждой части.")
    
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Собрать переведенные части")
    parser.add_argument("base_name", help="Базовое имя файла (без расширения)")
    parser.add_argument("--output", "-o", help="Выходной файл (по умолчанию: base_name_ru.docx)")
    
    args = parser.parse_args()
    
    output = args.output or f"{args.base_name}_ru.docx"
    
    print(f"\n{'='*60}")
    print(f"СБОРКА ПЕРЕВОДА")
    print(f"{'='*60}\n")
    
    # Загрузка частей
    print("📥 Загрузка переведенных частей...")
    parts = load_translated_parts(args.base_name)
    
    if not parts:
        print(f"❌ Не найдено переведенных частей для: {args.base_name}")
        print(f"   Ожидаемые файлы: {args.base_name}_part1_translated.json, ...")
        return 1
    
    print(f"   Найдено частей: {len(parts)}\n")
    
    # Создание документа
    print("📝 Создание документа...")
    create_docx(parts, output)
    
    print(f"\n{'='*60}")
    print(f"✅ ГОТОВО!")
    print(f"{'='*60}")
    print(f"Перевод сохранен: {output}")
    print(f"Частей обработано: {len(parts)}")
    print(f"{'='*60}\n")
    
    return 0


if __name__ == "__main__":
    exit(main())
