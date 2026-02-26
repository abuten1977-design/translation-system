#!/usr/bin/env python3
"""
Перевод с сохранением структуры документа
"""
import json
from pathlib import Path
from docx import Document
from copy import deepcopy


def translate_document_with_structure(input_docx, translations, output_docx):
    """Переводит документ с сохранением всей структуры"""
    # Загружаем оригинал
    doc = Document(input_docx)
    
    # Собираем весь переведенный текст
    translated_text = ""
    for trans_file in sorted(translations):
        with open(trans_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            translated_text += data['translated_text'] + "\n\n"
    
    # Разбиваем на параграфы
    translated_paras = [p.strip() for p in translated_text.split('\n\n') if p.strip()]
    
    # Заменяем текст в оригинальных параграфах
    para_index = 0
    for para in doc.paragraphs:
        if para.text.strip() and para_index < len(translated_paras):
            # Сохраняем форматирование, меняем только текст
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = translated_paras[para_index]
            else:
                para.text = translated_paras[para_index]
            para_index += 1
    
    # Переводим таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() and para_index < len(translated_paras):
                        for run in para.runs:
                            run.text = ''
                        if para.runs:
                            para.runs[0].text = translated_paras[para_index]
                        else:
                            para.text = translated_paras[para_index]
                        para_index += 1
    
    doc.save(output_docx)
    print(f"✅ Документ сохранен с сохранением структуры: {output_docx}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Использование: python3 translate_with_structure.py <base_name>")
        sys.exit(1)
    
    base_name = sys.argv[1]
    input_file = f"{base_name}.docx"
    output_file = f"{base_name}_ru_structured.docx"
    
    # Находим все файлы переводов
    translations = sorted(Path('.').glob(f"{base_name}_part*_translated.json"))
    
    if not translations:
        print(f"❌ Не найдено переводов для {base_name}")
        sys.exit(1)
    
    print(f"Найдено частей: {len(translations)}")
    translate_document_with_structure(input_file, translations, output_file)
