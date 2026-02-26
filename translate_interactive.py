#!/usr/bin/env python3
"""
Интерактивная система перевода через Kiro CLI
Не требует API ключа - все переводы выполняются через чат
"""
import argparse
import json
from pathlib import Path
from docx import Document


def extract_text_from_docx(filepath: str) -> str:
    """Извлекает текст из .docx файла"""
    doc = Document(filepath)
    return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def save_to_docx(text: str, output_path: str, metadata: dict = None):
    """Сохраняет текст в .docx файл"""
    doc = Document()
    doc.add_heading('Перевод', 0)
    
    # Добавляем переведенный текст
    for para in text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    # Добавляем метаданные если есть
    if metadata:
        doc.add_page_break()
        doc.add_heading('Информация о переводе', level=1)
        for key, value in metadata.items():
            doc.add_paragraph(f"{key}: {value}")
    
    doc.save(output_path)


def prepare_translation_request(text: str, source: str, target: str, mode: str) -> str:
    """Подготавливает запрос для перевода"""
    mode_desc = {
        "literary": "художественный перевод: передай смысл и тон максимально естественно и красиво",
        "precise": "точный перевод: передай смысл с максимальной точностью, используя правильные термины"
    }
    
    request = f"""
Выполни качественный перевод с {source} на {target}.
Режим: {mode_desc[mode]}

Используй 6-этапную проверку:
1. Начальный перевод
2. Проверка орфографии и грамматики
3. Проверка естественности
4. Проверка точности смысла
5. Проверка эмоционального тона
6. Финализация

Текст для перевода:
{text}

Верни результат в формате JSON:
{{
  "translated_text": "финальный перевод",
  "quality_report": {{
    "accuracy": "оценка точности",
    "fluency": "оценка естественности",
    "spelling": "оценка орфографии",
    "tone": "оценка тона"
  }},
  "notes": "краткие заметки о процессе"
}}
"""
    return request


def main():
    parser = argparse.ArgumentParser(
        description="Интерактивная система перевода (без API ключа)"
    )
    parser.add_argument("input", help="Входной файл (.docx)")
    parser.add_argument("output", help="Выходной файл (.docx)")
    parser.add_argument("--source", "-s", required=True, help="Язык оригинала")
    parser.add_argument("--target", "-t", required=True, help="Целевой язык")
    parser.add_argument(
        "--mode", "-m",
        choices=["literary", "precise"],
        default="precise",
        help="Режим перевода"
    )
    parser.add_argument(
        "--request-file", "-r",
        help="Сохранить запрос в файл вместо вывода"
    )
    
    args = parser.parse_args()
    
    # Проверка входного файла
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"❌ Ошибка: файл {args.input} не найден")
        return 1
    
    print(f"\n{'='*60}")
    print(f"ИНТЕРАКТИВНАЯ СИСТЕМА ПЕРЕВОДА")
    print(f"{'='*60}")
    print(f"Файл: {args.input}")
    print(f"Направление: {args.source} → {args.target}")
    print(f"Режим: {'Художественный' if args.mode == 'literary' else 'Точный'}")
    print(f"{'='*60}\n")
    
    # Извлечение текста
    print("📄 Извлечение текста из документа...")
    text = extract_text_from_docx(str(input_path))
    print(f"   Извлечено {len(text)} символов\n")
    
    # Подготовка запроса
    request = prepare_translation_request(text, args.source, args.target, args.mode)
    
    if args.request_file:
        # Сохранить запрос в файл
        with open(args.request_file, 'w', encoding='utf-8') as f:
            f.write(request)
        print(f"✅ Запрос сохранен в {args.request_file}")
        print(f"\nОтправь содержимое этого файла в чат для перевода.")
    else:
        # Вывести запрос
        print("="*60)
        print("ЗАПРОС ДЛЯ ПЕРЕВОДА")
        print("="*60)
        print("\nСкопируй текст ниже и отправь в чат:\n")
        print(request)
        print("\n" + "="*60)
        print("\nПосле получения ответа, сохрани JSON в файл и используй:")
        print(f"python3 save_translation.py <json-файл> {args.output}")
    
    return 0


if __name__ == "__main__":
    exit(main())
